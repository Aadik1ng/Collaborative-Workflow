"""
Script to generate assessment documents (Word and PDF).
Requires: pip install python-docx reportlab
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

# ============================================================================
# CONFIGURATION - UPDATE THESE WITH YOUR DETAILS
# ============================================================================
CONFIG = {
    "name": "Aaditya Aaryan",
    "email": "aadityaaryan639@gmail.com",
    "phone": "+91 8340118693",
    "date": datetime.now().strftime("%B %d, %Y"),
    "github_url": "https://github.com/Aadik1ng/Collaborative-Workflow",
    "vercel_url": "https://collaborative-workflow.vercel.app",
    "video_url": "https://your-video-link.com",
}

# ============================================================================
# DOCUMENT CONTENT
# ============================================================================

SECTIONS = {
    "project_overview": """
This project is a Real-Time Collaborative Workspace Backend - a production-ready API service that enables teams to collaborate on projects in real-time. It features secure authentication, project and workspace management, role-based access control, and asynchronous code execution jobs.

**Tech Stack:**
- **Framework**: FastAPI (Python 3.11)
- **Databases**: PostgreSQL (relational data), MongoDB (job results & logs), Redis (caching & pub/sub)
- **Authentication**: JWT with Argon2 password hashing
- **Async Workers**: Celery for background job processing
- **Real-Time**: WebSocket with Redis Pub/Sub

**Key Features:**
- User registration, login, and profile management
- Project CRUD with collaborator invitations
- Workspace management within projects
- Role-based access control (Owner, Collaborator, Viewer)
- Async code execution with status tracking
- Real-time collaboration via WebSockets
- Rate limiting and API caching
- Feature flags for runtime configuration
""",
    "architecture": """
The system follows a microservices-inspired architecture with clear separation of concerns:

**System Components:**
1. **FastAPI Application**: Handles HTTP/WebSocket requests, input validation, and routing
2. **PostgreSQL Database**: Stores users, projects, workspaces, and collaborator relationships
3. **MongoDB Database**: Stores high-velocity data like job results and activity logs
4. **Redis**: Serves as cache, rate limiter, pub/sub broker, and Celery message queue
5. **Celery Workers**: Process async jobs like code execution in the background

**Request Flow:**
1. Client sends request → Load Balancer → FastAPI instance
2. FastAPI validates input, checks auth (JWT), and rate limits (Redis)
3. Business logic executed, data fetched/stored in PostgreSQL/MongoDB
4. For async jobs: Task queued to Redis → Celery worker processes → Result stored in MongoDB
5. For real-time: WebSocket connections use Redis Pub/Sub for cross-instance messaging

**Database Schema (PostgreSQL):**
- `cw_users`: User accounts with hashed passwords
- `cw_projects`: Projects owned by users
- `cw_workspaces`: Workspaces within projects
- `cw_collaborators`: Many-to-many relationship with roles

**Key Libraries:**
- SQLAlchemy 2.0 with async support for PostgreSQL
- Motor for async MongoDB operations
- redis-py for async Redis operations
- passlib + argon2-cffi for password hashing
- PyJWT for token management
""",
    "setup_instructions": """
**Prerequisites:**
- Python 3.11+
- PostgreSQL, MongoDB, Redis (or use Docker)

**Local Setup:**

1. Clone the repository:
   ```
   git clone https://github.com/Aadik1ng/Collaborative-Workflow.git
   cd Collaborative-Workflow
   ```

2. Create virtual environment:
   ```
   python -m venv venv
   source venv/bin/activate  # Windows: venv\\Scripts\\activate
   ```

3. Install dependencies:
   ```
   pip install -r requirements.txt
   ```

4. Configure environment variables:
   ```
   cp .env.example .env
   # Edit .env with your database credentials
   ```

5. Start the API server:
   ```
   uvicorn app.main:app --reload
   ```

6. Start Celery worker (separate terminal):
   ```
   celery -A app.workers.celery_app worker --loglevel=info
   ```

**Docker Setup:**
```
cd docker
docker-compose up -d
```

**Environment Variables Required:**
- POSTGRES_URL: PostgreSQL connection string (with +asyncpg)
- MONGODB_URL: MongoDB connection string
- MONGODB_DATABASE: Database name
- REDIS_URL: Redis connection string
- SECRET_KEY: JWT signing secret
- ALGORITHM: JWT algorithm (HS256)
- ACCESS_TOKEN_EXPIRE_MINUTES: Token lifetime
- REFRESH_TOKEN_EXPIRE_DAYS: Refresh token lifetime
""",
    "design_decisions": """
**1. Dual Database Strategy (PostgreSQL + MongoDB)**
- Rationale: Relational data (users, projects, roles) benefits from ACID transactions. Non-relational data (job results, logs) needs flexible schemas and high write throughput.
- Trade-off: Increased operational complexity.

**2. Argon2 for Password Hashing**
- Rationale: Winner of Password Hashing Competition, resistant to GPU attacks, no 72-byte limit like bcrypt.
- Trade-off: Slightly higher CPU usage per hash.

**3. Celery for Async Jobs**
- Rationale: Decouples long-running tasks from request cycle, improves API responsiveness.
- Trade-off: Adds Redis as required dependency.

**4. JWT with Refresh Tokens**
- Rationale: Stateless access tokens enable horizontal scaling. Refresh tokens allow session invalidation.
- Trade-off: Requires careful token handling on client.

**5. Table Name Prefixing (cw_)**
- Rationale: Allows coexistence with other apps in shared database.
- Trade-off: Longer table names.

**6. Direct argon2-cffi Usage**
- Rationale: Bypasses passlib's backend detection issues in serverless environments.
- Trade-off: Less abstraction.
""",
    "scalability": """
**Horizontal Scaling:**
- FastAPI: Deploy multiple instances behind load balancer (stateless design)
- Celery Workers: Add more workers for increased job throughput
- PostgreSQL: Use read replicas, connection pooling (PgBouncer)
- MongoDB: Sharding for write scaling, replica sets for reads
- Redis: Redis Cluster for HA and scaling

**Performance Optimizations:**
- Sliding-window rate limiter protects against abuse
- Redis caching for frequently accessed data
- SQLAlchemy async connection pooling
- Idempotent job processing with unique IDs

**Database Indexing:**
- Indexed: user email, username, project owner_id, workspace project_id
- TTL indexes on activity logs (7-day expiry)

**Security Measures:**
- Argon2 password hashing
- JWT with short-lived access tokens
- Input validation with Pydantic
- CORS configuration
- Rate limiting per IP/user

**Future Enhancements:**
- Code execution sandboxing (Docker/gVisor)
- OpenTelemetry for distributed tracing
- Kubernetes deployment with Helm charts
""",
    "testing": """
**Running Tests:**

```
# Run all tests
pytest

# Run integration tests
pytest tests/integration

# Run unit tests
pytest tests/unit

# Run with coverage
pytest --cov=app --cov-report=term-missing

# Run with verbose output
pytest -v
```

**Test Coverage:**
- 58 tests covering authentication, projects, workspaces, collaborators, jobs
- ~59% code coverage focusing on critical paths
- Uses SQLite in-memory for fast integration tests

**Test Categories:**
- Unit Tests: Password hashing, JWT tokens, permissions
- Integration Tests: Full API request/response cycles with mocked databases
""",
    "deployment": """
**Vercel Deployment:**

1. Connect your GitHub repository to Vercel
2. Select "Other" as framework preset
3. Set environment variables in Vercel Dashboard:
   - POSTGRES_URL (with +asyncpg prefix)
   - MONGODB_URL
   - MONGODB_DATABASE
   - REDIS_URL
   - SECRET_KEY
   - Other JWT/app settings

4. Override Install Command: pip install -r requirements.txt
5. Deploy - Vercel auto-deploys on push to main

**Important Notes:**
- Vercel serverless has timeouts; use persistent hosting for WebSockets
- Celery workers must be deployed separately (Railway, Heroku, VPS)
- Use external databases (Railway, Atlas, Upstash)

**Docker Deployment:**
```
cd docker
docker-compose up -d --build
```

**Railway Deployment:**
1. Create PostgreSQL, MongoDB, Redis services
2. Deploy API from GitHub
3. Deploy Celery worker as separate service
4. Configure environment variables
""",
}

API_ENDPOINTS = [
    # Auth
    {
        "method": "POST",
        "route": "/api/v1/auth/register",
        "description": "Register new user",
        "request": '{"email": "user@example.com", "username": "user", "password": "SecurePass123!", "full_name": "John Doe"}',
        "response": '{"id": "uuid", "email": "...", "username": "...", "full_name": "..."}',
        "status": "201 Created",
    },
    {
        "method": "POST",
        "route": "/api/v1/auth/login",
        "description": "Login and get tokens",
        "request": '{"email": "user@example.com", "password": "SecurePass123!"}',
        "response": '{"access_token": "...", "refresh_token": "...", "token_type": "bearer"}',
        "status": "200 OK",
    },
    {
        "method": "POST",
        "route": "/api/v1/auth/refresh",
        "description": "Refresh access token",
        "request": '{"refresh_token": "..."}',
        "response": '{"access_token": "...", "token_type": "bearer"}',
        "status": "200 OK",
    },
    {
        "method": "POST",
        "route": "/api/v1/auth/logout",
        "description": "Logout user",
        "request": "Header: Authorization: Bearer <token>",
        "response": '{"message": "Logged out successfully"}',
        "status": "200 OK",
    },
    {
        "method": "GET",
        "route": "/api/v1/auth/me",
        "description": "Get current user",
        "request": "Header: Authorization: Bearer <token>",
        "response": '{"id": "...", "email": "...", "username": "...", "full_name": "..."}',
        "status": "200 OK",
    },
    {
        "method": "PUT",
        "route": "/api/v1/auth/me",
        "description": "Update profile",
        "request": '{"full_name": "New Name"}',
        "response": '{"id": "...", "email": "...", "full_name": "New Name"}',
        "status": "200 OK",
    },
    # Projects
    {
        "method": "POST",
        "route": "/api/v1/projects",
        "description": "Create project",
        "request": '{"name": "My Project", "description": "...", "is_public": false}',
        "response": '{"id": "uuid", "name": "...", "owner_id": "...", "created_at": "..."}',
        "status": "201 Created",
    },
    {
        "method": "GET",
        "route": "/api/v1/projects",
        "description": "List projects",
        "request": "Query: ?skip=0&limit=10",
        "response": '[{"id": "...", "name": "...", ...}]',
        "status": "200 OK",
    },
    {
        "method": "GET",
        "route": "/api/v1/projects/{id}",
        "description": "Get project",
        "request": "Path: project ID",
        "response": '{"id": "...", "name": "...", "owner": {...}, "workspaces": [...]}',
        "status": "200 OK",
    },
    {
        "method": "PUT",
        "route": "/api/v1/projects/{id}",
        "description": "Update project",
        "request": '{"name": "Updated Name"}',
        "response": '{"id": "...", "name": "Updated Name", ...}',
        "status": "200 OK",
    },
    {
        "method": "DELETE",
        "route": "/api/v1/projects/{id}",
        "description": "Delete project",
        "request": "Path: project ID",
        "response": '{"message": "Project deleted"}',
        "status": "200 OK",
    },
    # Workspaces
    {
        "method": "POST",
        "route": "/api/v1/projects/{id}/workspaces",
        "description": "Create workspace",
        "request": '{"name": "Workspace 1", "description": "..."}',
        "response": '{"id": "uuid", "name": "...", "project_id": "..."}',
        "status": "201 Created",
    },
    {
        "method": "GET",
        "route": "/api/v1/projects/{id}/workspaces",
        "description": "List workspaces",
        "request": "Path: project ID",
        "response": '[{"id": "...", "name": "...", ...}]',
        "status": "200 OK",
    },
    # Collaborators
    {
        "method": "POST",
        "route": "/api/v1/projects/{id}/collaborators",
        "description": "Invite collaborator",
        "request": '{"email": "collab@example.com", "role": "collaborator"}',
        "response": '{"id": "...", "user_id": "...", "role": "collaborator"}',
        "status": "201 Created",
    },
    {
        "method": "GET",
        "route": "/api/v1/projects/{id}/collaborators",
        "description": "List collaborators",
        "request": "Path: project ID",
        "response": '[{"user_id": "...", "email": "...", "role": "..."}]',
        "status": "200 OK",
    },
    # Jobs
    {
        "method": "POST",
        "route": "/api/v1/jobs",
        "description": "Submit code execution job",
        "request": '{"language": "python", "code": "print(\'Hello\')", "timeout": 30}',
        "response": '{"id": "uuid", "status": "pending", "created_at": "..."}',
        "status": "202 Accepted",
    },
    {
        "method": "GET",
        "route": "/api/v1/jobs/{id}",
        "description": "Get job status",
        "request": "Path: job ID",
        "response": '{"id": "...", "status": "completed", "output": "Hello", "execution_time": 0.5}',
        "status": "200 OK",
    },
    {
        "method": "GET",
        "route": "/api/v1/jobs",
        "description": "List user jobs",
        "request": "Query: ?skip=0&limit=10",
        "response": '[{"id": "...", "status": "...", ...}]',
        "status": "200 OK",
    },
    {
        "method": "POST",
        "route": "/api/v1/jobs/{id}/cancel",
        "description": "Cancel job",
        "request": "Path: job ID",
        "response": '{"message": "Job cancelled"}',
        "status": "200 OK",
    },
]


def create_word_document(output_path: str):
    """Create the Word document."""
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Title style
    title_style = styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 51, 102)

    # Heading 1 style modification
    h1 = styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor(0, 51, 102)

    # Heading 2 style modification
    h2 = styles["Heading 2"]
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor(0, 102, 153)

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Real-Time Collaborative Workspace Backend")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Backend Developer Assessment")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Candidate info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Name: {CONFIG['name']}\n").bold = True
    info_para.add_run(f"Email: {CONFIG['email']}\n")
    info_para.add_run(f"Phone: {CONFIG['phone']}\n")
    info_para.add_run(f"Date: {CONFIG['date']}")

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (Manual)
    # =========================================================================
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Project Overview",
        "2. Architecture Overview",
        "3. Setup & Run Instructions",
        "4. API Documentation",
        "5. Design Decisions & Trade-offs",
        "6. Scalability Considerations",
        "7. Testing Instructions",
        "8. Deployment Instructions",
        "9. Links Summary",
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1: PROJECT OVERVIEW
    # =========================================================================
    doc.add_heading("1. Project Overview", level=1)
    for line in SECTIONS["project_overview"].strip().split("\n"):
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 2: ARCHITECTURE
    # =========================================================================
    doc.add_heading("2. Architecture Overview", level=1)
    for line in SECTIONS["architecture"].strip().split("\n"):
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip().startswith(str(tuple(range(10)))):
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 3: SETUP INSTRUCTIONS
    # =========================================================================
    doc.add_heading("3. Setup & Run Instructions", level=1)
    in_code_block = False
    for line in SECTIONS["setup_instructions"].strip().split("\n"):
        if line.strip() == "```":
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: API DOCUMENTATION
    # =========================================================================
    doc.add_heading("4. API Documentation", level=1)

    current_category = None
    for endpoint in API_ENDPOINTS:
        route = endpoint["route"]
        if "/auth/" in route and current_category != "Authentication":
            current_category = "Authentication"
            doc.add_heading("Authentication Endpoints", level=2)
        elif (
            "/projects" in route
            and "/workspaces" not in route
            and "/collaborators" not in route
            and current_category != "Projects"
        ):
            current_category = "Projects"
            doc.add_heading("Project Endpoints", level=2)
        elif "/workspaces" in route and current_category != "Workspaces":
            current_category = "Workspaces"
            doc.add_heading("Workspace Endpoints", level=2)
        elif "/collaborators" in route and current_category != "Collaborators":
            current_category = "Collaborators"
            doc.add_heading("Collaborator Endpoints", level=2)
        elif "/jobs" in route and current_category != "Jobs":
            current_category = "Jobs"
            doc.add_heading("Job Endpoints", level=2)

        # Endpoint header
        p = doc.add_paragraph()
        run = p.add_run(f"{endpoint['method']} {endpoint['route']}")
        run.font.bold = True
        run.font.size = Pt(11)

        doc.add_paragraph(f"Description: {endpoint['description']}")

        # Request
        p = doc.add_paragraph()
        p.add_run("Request: ").bold = True
        req_run = p.add_run(endpoint["request"])
        req_run.font.name = "Courier New"
        req_run.font.size = Pt(9)

        # Response
        p = doc.add_paragraph()
        p.add_run("Response: ").bold = True
        res_run = p.add_run(endpoint["response"])
        res_run.font.name = "Courier New"
        res_run.font.size = Pt(9)

        doc.add_paragraph(f"Status: {endpoint['status']}")
        doc.add_paragraph()  # Spacer

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: DESIGN DECISIONS
    # =========================================================================
    doc.add_heading("5. Design Decisions & Trade-offs", level=1)
    for line in SECTIONS["design_decisions"].strip().split("\n"):
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: SCALABILITY
    # =========================================================================
    doc.add_heading("6. Scalability Considerations", level=1)
    for line in SECTIONS["scalability"].strip().split("\n"):
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: TESTING
    # =========================================================================
    doc.add_heading("7. Testing Instructions", level=1)
    in_code_block = False
    for line in SECTIONS["testing"].strip().split("\n"):
        if line.strip() == "```":
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 8: DEPLOYMENT
    # =========================================================================
    doc.add_heading("8. Deployment Instructions", level=1)
    in_code_block = False
    for line in SECTIONS["deployment"].strip().split("\n"):
        if line.strip() == "```":
            in_code_block = not in_code_block
            continue
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.add_page_break()

    # =========================================================================
    # SECTION 9: LINKS SUMMARY
    # =========================================================================
    doc.add_heading("9. Links Summary", level=1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("GitHub Repository: ").bold = True
    p.add_run(CONFIG["github_url"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Live Vercel Deployment: ").bold = True
    p.add_run(CONFIG["vercel_url"])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Walkthrough Video: ").bold = True
    p.add_run(CONFIG["video_url"])

    # Save the document
    doc.save(output_path)
    print(f"✅ Word document created: {output_path}")


def create_pdf_document(output_path: str):
    """Create the PDF document using ReportLab."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    # Styles
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=24,
        textColor=colors.HexColor("#003366"),
        spaceAfter=20,
        alignment=1,  # Center
    )

    h1_style = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontSize=16,
        textColor=colors.HexColor("#003366"),
        spaceBefore=20,
        spaceAfter=10,
    )

    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#006699"),
        spaceBefore=15,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        "CustomBody", parent=styles["Normal"], fontSize=10, spaceAfter=6, leading=14
    )

    code_style = ParagraphStyle(
        "CustomCode",
        parent=styles["Code"],
        fontSize=9,
        fontName="Courier",
        backColor=colors.HexColor("#f5f5f5"),
        leftIndent=10,
        spaceAfter=4,
    )

    story = []

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    story.append(Spacer(1, 2 * inch))
    story.append(Paragraph("Real-Time Collaborative Workspace Backend", title_style))
    story.append(Spacer(1, 0.3 * inch))
    story.append(
        Paragraph(
            "Backend Developer Assessment",
            ParagraphStyle(
                "Subtitle", parent=styles["Normal"], fontSize=14, textColor=colors.gray, alignment=1
            ),
        )
    )
    story.append(Spacer(1, 1 * inch))

    cover_info = f"""
    <b>Name:</b> {CONFIG["name"]}<br/>
    <b>Email:</b> {CONFIG["email"]}<br/>
    <b>Phone:</b> {CONFIG["phone"]}<br/>
    <b>Date:</b> {CONFIG["date"]}
    """
    story.append(
        Paragraph(
            cover_info,
            ParagraphStyle(
                "CoverInfo", parent=styles["Normal"], fontSize=12, alignment=1, leading=18
            ),
        )
    )
    story.append(PageBreak())

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    story.append(Paragraph("Table of Contents", h1_style))
    toc_items = [
        "1. Project Overview",
        "2. Architecture Overview",
        "3. Setup & Run Instructions",
        "4. API Documentation",
        "5. Design Decisions & Trade-offs",
        "6. Scalability Considerations",
        "7. Testing Instructions",
        "8. Deployment Instructions",
        "9. Links Summary",
    ]
    for item in toc_items:
        story.append(Paragraph(item, body_style))
    story.append(PageBreak())

    # =========================================================================
    # SECTIONS
    # =========================================================================

    def add_section(title, content):
        story.append(Paragraph(title, h1_style))
        in_code_block = False
        for line in content.strip().split("\n"):
            line = line.replace("<", "&lt;").replace(">", "&gt;")
            if line.strip() == "```":
                in_code_block = not in_code_block
                continue
            if in_code_block:
                story.append(Paragraph(line or " ", code_style))
            elif line.startswith("**") and line.endswith("**"):
                story.append(Paragraph(f"<b>{line.strip('*')}</b>", body_style))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            elif line.strip():
                story.append(Paragraph(line, body_style))
        story.append(PageBreak())

    add_section("1. Project Overview", SECTIONS["project_overview"])
    add_section("2. Architecture Overview", SECTIONS["architecture"])
    add_section("3. Setup & Run Instructions", SECTIONS["setup_instructions"])

    # API Documentation (special handling)
    story.append(Paragraph("4. API Documentation", h1_style))

    current_category = None
    for endpoint in API_ENDPOINTS:
        route = endpoint["route"]
        if "/auth/" in route and current_category != "Auth":
            current_category = "Auth"
            story.append(Paragraph("Authentication Endpoints", h2_style))
        elif (
            "/projects" in route
            and "/workspaces" not in route
            and "/collaborators" not in route
            and current_category != "Projects"
        ):
            current_category = "Projects"
            story.append(Paragraph("Project Endpoints", h2_style))
        elif "/workspaces" in route and current_category != "Workspaces":
            current_category = "Workspaces"
            story.append(Paragraph("Workspace Endpoints", h2_style))
        elif "/collaborators" in route and current_category != "Collaborators":
            current_category = "Collaborators"
            story.append(Paragraph("Collaborator Endpoints", h2_style))
        elif "/jobs" in route and current_category != "Jobs":
            current_category = "Jobs"
            story.append(Paragraph("Job Endpoints", h2_style))

        story.append(Paragraph(f"<b>{endpoint['method']} {endpoint['route']}</b>", body_style))
        story.append(Paragraph(f"Description: {endpoint['description']}", body_style))
        req = endpoint["request"].replace("<", "&lt;").replace(">", "&gt;")
        res = endpoint["response"].replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(f"Request: <font face='Courier' size='9'>{req}</font>", body_style))
        story.append(Paragraph(f"Response: <font face='Courier' size='9'>{res}</font>", body_style))
        story.append(Paragraph(f"Status: {endpoint['status']}", body_style))
        story.append(Spacer(1, 0.1 * inch))

    story.append(PageBreak())

    add_section("5. Design Decisions & Trade-offs", SECTIONS["design_decisions"])
    add_section("6. Scalability Considerations", SECTIONS["scalability"])
    add_section("7. Testing Instructions", SECTIONS["testing"])
    add_section("8. Deployment Instructions", SECTIONS["deployment"])

    # Links Summary
    story.append(Paragraph("9. Links Summary", h1_style))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"<b>GitHub Repository:</b> {CONFIG['github_url']}", body_style))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(f"<b>Live Vercel Deployment:</b> {CONFIG['vercel_url']}", body_style))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph(f"<b>Walkthrough Video:</b> {CONFIG['video_url']}", body_style))

    # Build PDF
    doc.build(story)
    print(f"✅ PDF document created: {output_path}")


def main():
    """Generate both documents."""
    # Output directory
    output_dir = os.path.dirname(os.path.abspath(__file__))

    # Generate Word document
    docx_path = os.path.join(output_dir, "Backend_Assessment_Submission.docx")
    create_word_document(docx_path)

    # Generate PDF document
    pdf_path = os.path.join(output_dir, "Backend_Assessment_Submission.pdf")
    create_pdf_document(pdf_path)

    print("\n" + "=" * 60)
    print("📄 Documents generated successfully!")
    print("=" * 60)
    print(f"\n📝 Word: {docx_path}")
    print(f"📕 PDF:  {pdf_path}")
    print("\n⚠️  Remember to update the CONFIG section at the top of this")
    print("    script with your personal details before final submission!")


if __name__ == "__main__":
    main()
